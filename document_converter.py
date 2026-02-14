from __future__ import annotations

import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Literal, cast, Any

from docx import Document
from docx.shared import Inches, Mm, Pt, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK, WD_LINE_SPACING
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from docx.styles.style import ParagraphStyle
from docx.shared import Inches, Mm, Pt, Emu, RGBColor

_HAS_NB = False
try:
    import nbformat as nbf  # type: ignore[import-untyped]
    from nbconvert import HTMLExporter, MarkdownExporter  # type: ignore[import-untyped]
    from traitlets.config import Config  # type: ignore[import-untyped]
    _HAS_NB = True
except ImportError:
    pass

if TYPE_CHECKING:
    import nbformat as nbf  # type: ignore[no-redef]
    from nbconvert import HTMLExporter, MarkdownExporter  # type: ignore[no-redef]
    from traitlets.config import Config  # type: ignore[no-redef]


class ConversionError(Exception):
    """Raised when a conversion step fails."""


@dataclass
class ConversionResult:
    """Paths to generated output files."""
    docx: Path | None = None
    html: Path | None = None
    markdown: Path | None = None


@dataclass
class DocumentConverter:
    """Converts Jupyter notebooks or Markdown files to DOCX.
    
    Args:
        template: Path to a reference DOCX template for styling.
        hide_code: If True, exclude code cell inputs from notebook exports.
        keep_text: If True, preserve text outputs from code cells.
        export_html: If True, also export notebooks to HTML.
        export_markdown: If True, also keep the intermediate Markdown export.
        include_toc: If True, adds a Table of Contents at the beginning.
        doc_title: Title to use in the DOCX header (defaults to filename).
        paper_size: 'A4' or 'Letter'. Defaults to 'A4'.
    """

    template: Path | None = None
    hide_code: bool = False
    keep_text: bool = False
    export_html: bool = False
    export_markdown: bool = False
    include_toc: bool = False
    doc_title: str | None = None
    paper_size: Literal["A4", "Letter"] = "A4"
    header_text: str | None = None
    page_number_pos: str = "right"
    show_page_word: bool = False
    text_align: str = "justify"
    font_family: str = "Aptos"
    font_size_body: int = 12
    font_size_table: int = 11
    font_size_header: int = 9
    font_size_code: int = 10
    resize_images: bool = True
    resize_tables: bool = True

    def __post_init__(self) -> None:
        if self.template is not None:
            self.template = Path(self.template).resolve()


    def convert(self, input_path: str | Path, output_dir: str | Path | None = None, base_name: str | None = None) -> ConversionResult:
        input_path = Path(input_path).resolve()
        if not input_path.exists(): raise ConversionError(f"Input file not found: {input_path}")
        suffix = input_path.suffix.lower()
        if suffix not in (".ipynb", ".md"): raise ConversionError(f"Unsupported file type: {suffix}")
        if output_dir is None: output_dir = input_path.parent
        else: output_dir = Path(output_dir).resolve(); output_dir.mkdir(parents=True, exist_ok=True)
        if base_name is None: base_name = input_path.stem
        if suffix == ".ipynb": return self._process_notebook(input_path, base_name, output_dir)
        else: return self._process_markdown(input_path, base_name, output_dir)

    @staticmethod
    def _require_nb() -> None:
        if not _HAS_NB: raise ConversionError("nbformat/nbconvert not available.")

    @staticmethod
    def _require_pandoc() -> None:
        if shutil.which("pandoc") is None: raise ConversionError("pandoc not found on PATH.")

    @staticmethod
    def _clean_outputs(nb: nbf.NotebookNode, keep_text: bool = False) -> nbf.NotebookNode:
        for cell in nb.cells:
            if cell.get("cell_type") != "code": continue
            if not keep_text:
                new_outputs = []
                for out in cell.get("outputs", []):
                    otype = out.get("output_type"); data = out.get("data", {})
                    has_img = any(k in data for k in ("image/png", "image/svg+xml"))
                    has_html = "text/html" in data
                    if otype in {"display_data", "execute_result"} and (has_img or has_html): new_outputs.append(out)
                cell["outputs"] = new_outputs
            cell["execution_count"] = None
        return nb
    
    def _export_html(self, nb: nbf.NotebookNode, out_path: Path) -> None:
        self._require_nb()
        c = Config(); c.HTMLExporter.exclude_input_prompt = True; c.HTMLExporter.exclude_output_prompt = True
        if self.hide_code: c.HTMLExporter.exclude_input = True
        html_exporter = HTMLExporter(config=c)
        body, _ = html_exporter.from_notebook_node(nb)
        out_path.write_text(body, encoding="utf-8")

    @staticmethod
    def _html_table_to_markdown(html: str) -> str:
        """Convert an HTML <table> to a Markdown pipe table."""
        from lxml.html import fromstring

        doc = fromstring(html)
        table = doc if doc.tag == 'table' else doc.find('.//table')
        if table is None:
            return ''

        rows = []
        for tr in table.iter('tr'):
            cells = [c.text_content().strip() for c in tr if c.tag in ('td', 'th')]
            if cells:
                rows.append(cells)

        if not rows:
            return ''

        num_cols = max(len(r) for r in rows)
        for r in rows:
            while len(r) < num_cols:
                r.append('')

        lines = []
        lines.append('| ' + ' | '.join(rows[0]) + ' |')
        lines.append('| ' + ' | '.join(['---'] * num_cols) + ' |')
        for row in rows[1:]:
            lines.append('| ' + ' | '.join(row) + ' |')

        return '\n'.join(lines)

    def _strip_captions(self, md_content: str) -> str:
        return re.sub(r'!\[.*?\]\((.*?)\)', r'![](\1)', md_content)        

    def _export_markdown(self, nb: nbf.NotebookNode, out_path: Path, assets_dir: str) -> None:
        self._require_nb()
        c = Config(); c.MarkdownExporter.exclude_input_prompt = True; c.MarkdownExporter.exclude_output_prompt = True
        if self.hide_code: c.MarkdownExporter.exclude_input = True
        resources = {"output_files_dir": assets_dir}
        md_exporter = MarkdownExporter(config=c)
        body, resources = md_exporter.from_notebook_node(nb, resources=resources)
        
        # Strip captions
        body = self._strip_captions(body)
        body = self._strip_hrs_after_headers(body)
        
        out_path.write_text(body, encoding="utf-8")
        for fname, data in resources.get("outputs", {}).items():
            img_path = out_path.parent / fname; img_path.parent.mkdir(parents=True, exist_ok=True)
            img_path.write_bytes(data)

    def _set_section_properties(self, section, title: str):
        """Helper to set margins and page size for a section."""
        # 1. Margins (in inches)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.8)

        # 2. Paper Size
        if self.paper_size == "A4":
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        else: # Letter
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        # 3. Header/Footer Distances (Programmatic XML update)
        # Word stores these in the w:pgMar element as Twips (1/20th of a point).
        # python-docx Section object doesn't expose these directly.
        try:
            sectPr = section._sectPr
            pgMar = sectPr.find(qn('w:pgMar'))
            
            if pgMar is not None:
                # Calculate Twips: Inches -> Points -> Twips (Pt * 20)
                # Header distance: 0.3 inches
                header_dist_twips = int(Inches(0.3).pt * 20)
                # Footer distance: 0.2 inches
                footer_dist_twips = int(Inches(0.2).pt * 20)
                
                # Set attributes in XML
                pgMar.set(qn('w:header'), str(header_dist_twips))
                pgMar.set(qn('w:footer'), str(footer_dist_twips))
        except Exception as e:
            print(f"Warning: Could not set header/footer distances: {e}")

    def _strip_hrs_after_headers(self, md_content: str) -> str:
        """Removes Horizontal Rules appearing immediately after Header 1."""
        lines = md_content.split('\n')
        new_lines = []
        
        def is_hr(text):
            t = text.strip()
            if t in ['---', '***', '___']:
                return True
            # Check for long lines of underscores or hyphens
            if len(t) > 5:
                chars = set(t)
                if chars == {'_'} or chars == {'-'} or chars == {'*'}:
                    return True
            return False

        def is_header_1(text):
            t = text.strip()
            # Must start with '#'
            if not t.startswith('#'):
                return False
            # Count hashes
            count = 0
            for char in t:
                if char == '#':
                    count += 1
                else:
                    break
            # Must be exactly 1 hash, followed by a space
            if count == 1 and len(t) > 1 and t[1] == ' ':
                return True
            return False

        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Check if CURRENT line is an HR
            if is_hr(line):
                # Look BACKWARDS to find previous non-empty line
                prev_idx = i - 1
                while prev_idx >= 0 and lines[prev_idx].strip() == "":
                    prev_idx -= 1
                
                is_after_h1 = False
                if prev_idx >= 0:
                    prev_line = lines[prev_idx].strip()
                    if is_header_1(prev_line):
                        is_after_h1 = True
                
                # Only keep the HR if it is NOT after H1
                if not is_after_h1:
                    new_lines.append(line)
            else:
                new_lines.append(line)
            
            i += 1
            
        return '\n'.join(new_lines)
    
    def _style_captions_in_markdown(self, md_content: str) -> str:
        """Wraps Figure captions in span tags for 10pt font size."""
        lines = md_content.split('\n')
        new_lines = []
        
        # Regex to find "Figure " followed by a number
        # e.g., "Figure 1: Description"
        fig_pattern = re.compile(r'^(Figure\s+\d+.*)$', re.IGNORECASE)
        
        for line in lines:
            match = fig_pattern.match(line.strip())
            if match:
                # Wrap the entire line in a span with font-size: 10pt
                # We preserve the original line content inside the span
                new_lines.append(f'<span style="font-size:10pt">{line}</span>')
            else:
                new_lines.append(line)
        
        return '\n'.join(new_lines)    

    def _apply_styling(self, docx_path: Path) -> None:
        """Apply global font styles, justification, and resizing."""
        doc = Document(str(docx_path))
        
        # Define Colors
        accent_blue = RGBColor(0x4F, 0x81, 0xBD)

        # 1. Setup Base Style (Defaults)
        style = doc.styles['Normal']
        p_style = cast(ParagraphStyle, style)
        font = p_style.font
        font.name = self.font_family
        font.size = Pt(self.font_size_body)

        # 2. Handle TOC Cover Page
        if self.include_toc and len(doc.paragraphs) > 0:
            cover_para = doc.add_paragraph("INSERT YOUR COVER PAGE HERE")
            cover_para.style = 'Normal'
            run = cover_para.add_run()
            run.add_break(WD_BREAK.PAGE)
            doc.element.body.insert(0, cover_para._element)

        # 3. PASS 1: Cleanup (Iterate BACKWARDS to safely remove elements)
        # We use a list() copy to avoid issues with the live proxy
        paragraphs = list(doc.paragraphs)
        for i in range(len(paragraphs) - 1, -1, -1):
            paragraph = paragraphs[i]
            style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
            text = paragraph.text.strip()
            
            # --- HR Removal Logic ---
            is_hr = False
            # Check Borders
            pPr = paragraph._p.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:pBdr')) is not None:
                is_hr = True
            # Check Text (---, ___, etc)
            if text in ['---', '***', '___'] or (len(text) > 5 and set(text) == {'_'}):
                is_hr = True
            
            if is_hr:
                # Look backwards to find previous non-empty paragraph
                prev_idx = i - 1
                while prev_idx >= 0 and paragraphs[prev_idx].text.strip() == "":
                    prev_idx -= 1
                
                if prev_idx >= 0:
                    _style = paragraphs[prev_idx].style
                    prev_style = _style.name if _style and _style.name else ""

                    if prev_style == 'Heading 1':
                        # Remove it
                        paragraph._element.getparent().remove(paragraph._element)
                        # Important: Remove from our local list too so next step doesn't fail
                        # Actually, we are done with this pass, so not strictly necessary for logic
                        # but good for hygiene.
                        pass

        # 4. PASS 2: Styling (Iterate FORWARDS)
        # Refresh paragraph list because we might have deleted some
        paragraphs = list(doc.paragraphs)
        
        for i, paragraph in enumerate(paragraphs):
            style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
            text = paragraph.text.strip()
            
            is_code = 'Source Code' in style_name
            is_heading = style_name.startswith('Heading')
            pPr = paragraph._p.find(qn('w:pPr'))
            has_numPr = pPr is not None and pPr.find(qn('w:numPr')) is not None
            is_list = 'List' in style_name or 'Compact' in style_name or has_numPr
            
            # --- A. Body Text Spacing ---
            has_image = len(paragraph._element.xpath('.//w:drawing')) > 0 or \
                        len(paragraph._element.xpath('.//w:pict')) > 0
            _math_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
            has_math = len(paragraph._element.findall(f'.//{{{_math_ns}}}oMath')) > 0 or \
                       len(paragraph._element.findall(f'.//{{{_math_ns}}}oMathPara')) > 0
            paragraph.paragraph_format.space_before = Pt(9)
            paragraph.paragraph_format.space_after = Pt(9)
            if not has_image and not has_math:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(16)

            is_caption = False  # Reset each iteration

            # --- B. Caption Styling (Fix Issue 2) ---
            # Check if this paragraph looks like a caption
            # Regex: Starts with "Figure " followed by a digit
            if re.match(r'^Figure\s+\d+', text, re.IGNORECASE):






                # Check if previous paragraph contains an Image
                prev_para = paragraphs[i-1] if i > 0 else None
                if prev_para:
                    # Check for image XML in previous paragraph
                    prev_has_image = len(prev_para._element.xpath('.//w:drawing')) > 0 or \
                                     len(prev_para._element.xpath('.//w:pict')) > 0

                    if prev_has_image:
                        # Apply Caption Style: 10pt font, 13pt space before
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = self.font_family
                        
                        # Override spacing for captions
                        paragraph.paragraph_format.space_before = Pt(9)
                        paragraph.paragraph_format.space_after = Pt(9)
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        is_caption = True
            
            # --- C. Heading & Code Logic ---
            if is_heading:
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.page_break_before = (style_name == 'Heading 1')

            # --- TOC Heading: Style like Heading 1 ---
            for sdt in doc.element.body.iterchildren(qn('w:sdt')):
                sdtContent = sdt.find(qn('w:sdtContent'))
                if sdtContent is not None:
                    for p in sdtContent.iterchildren(qn('w:p')):
                        pPr = p.find(qn('w:pPr'))
                        if pPr is not None:
                            pStyle = pPr.find(qn('w:pStyle'))
                            if pStyle is not None and pStyle.get(qn('w:val')) == 'TOCHeading':
                                # Set 6pt space after
                                spacing = pPr.find(qn('w:spacing'))
                                if spacing is None:
                                    spacing = OxmlElement('w:spacing')
                                    pPr.append(spacing)
                                spacing.set(qn('w:after'), '120')  # 6pt = 120 twips
                                for r in p.iter(qn('w:r')):
                                    rPr = r.find(qn('w:rPr'))
                                    if rPr is None:
                                        rPr = OxmlElement('w:rPr')
                                        r.insert(0, rPr)
                                    sz = rPr.find(qn('w:sz'))
                                    if sz is None:
                                        sz = OxmlElement('w:sz')
                                        rPr.append(sz)
                                    sz.set(qn('w:val'), '36')  # 18pt = 36 half-points
                                    color = rPr.find(qn('w:color'))
                                    if color is None:
                                        color = OxmlElement('w:color')
                                        rPr.append(color)
                                    color.set(qn('w:val'), '4F81BD')  # accent_blue

            if is_code:
                def _adj_is_code(idx):
                    if idx < 0 or idx >= len(paragraphs):
                        return False
                    s = paragraphs[idx].style
                    return 'Source Code' in (s.name if s and s.name else '')
                is_first_code = not _adj_is_code(i - 1)
                is_last_code = not _adj_is_code(i + 1)

                # Spacing compensates for border space extending the grey area
                paragraph.paragraph_format.space_before = Pt(12) if is_first_code else Pt(0)
                paragraph.paragraph_format.space_after = Pt(9) if is_last_code else Pt(0)

                # Light grey background
                pPr_el = paragraph._p.get_or_add_pPr()
                for old_shd in pPr_el.findall(qn('w:shd')):
                    pPr_el.remove(old_shd)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F2F2')
                pPr_el.append(shd)

                # Padding via invisible borders with space:
                # Left/right on all code paragraphs, top/bottom on first/last only
                for old_bdr in pPr_el.findall(qn('w:pBdr')):
                    pPr_el.remove(old_bdr)
                pBdr = OxmlElement('w:pBdr')
                for side in ('left', 'right'):
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '9')  # 9pt text inset
                    border.set(qn('w:color'), 'auto')
                    pBdr.append(border)
                if is_first_code:
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'none')
                    top.set(qn('w:sz'), '0')
                    top.set(qn('w:space'), '5')
                    top.set(qn('w:color'), 'auto')
                    pBdr.append(top)
                if is_last_code:
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'none')
                    bottom.set(qn('w:sz'), '0')
                    bottom.set(qn('w:space'), '9')
                    bottom.set(qn('w:color'), 'auto')
                    pBdr.append(bottom)
                pPr_el.append(pBdr)

                # Indent offsets left/right border space so grey aligns with margins
                for old_ind in pPr_el.findall(qn('w:ind')):
                    pPr_el.remove(old_ind)
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), '180')   # 9pt = 180tw
                ind.set(qn('w:right'), '180')
                pPr_el.append(ind)

            # --- List Styling ---
            if is_list:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph.paragraph_format.line_spacing = 1.16

                # Remove paragraph-level w:ind — let numbering definition control indent
                pPr_el = paragraph._p.get_or_add_pPr()
                for existing_ind in pPr_el.findall(qn('w:ind')):
                    pPr_el.remove(existing_ind)
                # Also remove any tab stops
                for tabs in pPr_el.findall(qn('w:tabs')):
                    pPr_el.remove(tabs)

            # --- D. Alignment & Font ---
            if is_code:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif has_image:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                if self.text_align == "justify":
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            for run in paragraph.runs:
                if not is_code:
                    run.font.name = self.font_family
                
                if is_code:
                    run.font.size = Pt(self.font_size_code)
                elif is_heading:
                    if style_name == "Heading 1":
                        run.font.size = Pt(18)
                elif not is_caption:
                    run.font.size = Pt(self.font_size_body)
                
                if is_heading:
                    run.font.color.rgb = accent_blue

        # 5. Set numbering definition indents (controls visual list indent)
        # Progressive indentation per level:
        #   Level 0: Left=0.45" (648tw), Hanging=0.20" (288tw) → bullet at 0.25", text at 0.45"
        #   Level 1: Left=0.70" (1008tw), Hanging=0.20" (288tw) → bullet at 0.50", text at 0.70"
        #   Level 2: Left=0.95" (1368tw), Hanging=0.20" (288tw) → bullet at 0.75", text at 0.95"
        #   etc.
        try:
            numbering_part = doc.part.numbering_part
            numbering_xml = numbering_part._element
            for abstractNum in numbering_xml.findall(qn('w:abstractNum')):
                for lvl in abstractNum.findall(qn('w:lvl')):
                    ilvl = int(lvl.get(qn('w:ilvl'), '0'))
                    left_twips = str(648 + ilvl * 360)  # 0.45" + level * 0.25"
                    lvl_pPr = lvl.find(qn('w:pPr'))
                    if lvl_pPr is None:
                        lvl_pPr = OxmlElement('w:pPr')
                        lvl.append(lvl_pPr)
                    ind = lvl_pPr.find(qn('w:ind'))
                    if ind is None:
                        ind = OxmlElement('w:ind')
                        lvl_pPr.append(ind)
                    ind.set(qn('w:left'), left_twips)
                    ind.set(qn('w:hanging'), '288')   # 0.20" hanging for all levels

                    # Set white bullet "○" for level 1 (sub-items)
                    if ilvl == 1:
                        numFmt = lvl.find(qn('w:numFmt'))
                        if numFmt is not None and numFmt.get(qn('w:val')) == 'bullet':
                            lvlText = lvl.find(qn('w:lvlText'))
                            if lvlText is not None:
                                lvlText.set(qn('w:val'), 'o')
                            # Set font to Courier New for the outline bullet
                            lvl_rPr = lvl.find(qn('w:rPr'))
                            if lvl_rPr is None:
                                lvl_rPr = OxmlElement('w:rPr')
                                lvl.append(lvl_rPr)
                            for old_fonts in lvl_rPr.findall(qn('w:rFonts')):
                                lvl_rPr.remove(old_fonts)
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:ascii'), 'Courier New')
                            rFonts.set(qn('w:hAnsi'), 'Courier New')
                            rFonts.set(qn('w:cs'), 'Courier New')
                            rFonts.set(qn('w:hint'), 'default')
                            lvl_rPr.append(rFonts)
        except Exception:
            pass  # No numbering part = no lists

        # 6. Configure Tables
        section = doc.sections[0]
        
        page_w = section.page_width if section.page_width else Inches(8.5)
        left_m = section.left_margin if section.left_margin else Inches(1.0)
        right_m = section.right_margin if section.right_margin else Inches(0.8)
        available_width = page_w - left_m - right_m

        for table in doc.tables:
            # Apply Table Grid as clean base style
            try:
                if 'Table Grid' in doc.styles:
                    table.style = 'Table Grid'
            except Exception:
                pass

            tblPr = table._tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                table._tbl.insert(0, tblPr)

            # Clear all Pandoc table-level overrides
            for tag in ('w:tblBorders', 'w:tblCellMar', 'w:tblLook'):
                for el in tblPr.findall(qn(tag)):
                    tblPr.remove(el)

            # Set table-level borders (thin single #156082)
            tblBorders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '4F81BD')
                tblBorders.append(border)
            tblPr.append(tblBorders)

            # Set table-level default cell margins
            tblCellMar = OxmlElement('w:tblCellMar')
            for side, val in [('top', '58'), ('bottom', '58'), ('left', '115'), ('right', '115')]:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), val)
                m.set(qn('w:type'), 'dxa')
                tblCellMar.append(m)
            tblPr.append(tblCellMar)

            # Set tblLook
            tblLook = OxmlElement('w:tblLook')
            tblLook.set(qn('w:val'), '04A0')
            tblLook.set(qn('w:firstRow'), '1')
            tblLook.set(qn('w:lastRow'), '0')
            tblLook.set(qn('w:firstColumn'), '1')
            tblLook.set(qn('w:lastColumn'), '0')
            tblLook.set(qn('w:noHBand'), '0')
            tblLook.set(qn('w:noVBand'), '1')
            tblPr.append(tblLook)

            # Center table horizontally
            for old_jc in tblPr.findall(qn('w:jc')):
                tblPr.remove(old_jc)
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            tblPr.append(jc)

            # Apply per-cell formatting
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    # Clear all Pandoc cell overrides
                    for tag in ('w:shd', 'w:tcBorders', 'w:tcMar'):
                        for el in tcPr.findall(qn(tag)):
                            tcPr.remove(el)

                    # Vertical alignment: center
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'center')
                    tcPr.append(vAlign)

                    # Header row: blue background
                    if i == 0:
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), '4F81BD')
                        tcPr.append(shd)

                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        pPr = paragraph._p.find(qn('w:pPr'))
                        if pPr is not None:
                            for shd in pPr.findall(qn('w:shd')):
                                pPr.remove(shd)
                        for run in paragraph.runs:
                            run.font.size = Pt(self.font_size_table)
                            run.font.name = self.font_family
                            if i == 0:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            else:
                                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            if self.resize_tables:
                table.autofit = False
                cast(Any, table).preferred_width = available_width

        # 6. Resize Images (Same as before)
        if self.resize_images:
            for shape in doc.inline_shapes:
                s = cast(Any, shape)
                current_width = Emu(s.width)
                target_width = Emu(available_width)
                
                if current_width != target_width:
                    ratio = target_width / current_width
                    s.width = target_width
                    s.height = Emu(s.height * ratio)

        doc.save(str(docx_path))

    def _add_header_footer(self, docx_path: Path, title: str) -> None:
        """Add header, footer, margins, and page layout."""
        doc = Document(str(docx_path))

        header_content = self.header_text if self.header_text else title

        for section in doc.sections:
            self._set_section_properties(section, title)
            section.different_first_page_header_footer = True

            # --- HEADER ---
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_content
            
            for run in header_para.runs:
                run.font.size = Pt(self.font_size_header)
                run.font.name = self.font_family

            # Line BELOW header (Black, with 2pt spacing)
            self._set_paragraph_border(header_para, border_type="bottom", color="000000", size="6", space="2")

            # --- FOOTER ---
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            
            for run in footer_para.runs:
                run.font.size = Pt(self.font_size_header)
                run.font.name = self.font_family
            
            # Alignment
            if self.page_number_pos == 'left':
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif self.page_number_pos == 'right':
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            if self.show_page_word:
                run = footer_para.add_run("Page ")
                run.font.size = Pt(self.font_size_header)
                run.font.name = self.font_family

            run = footer_para.add_run()
            run.font.size = Pt(self.font_size_header)
            run.font.name = self.font_family
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

            # Line ABOVE footer (Black)
            self._set_paragraph_border(footer_para, border_type="top", color="000000", size="6", space="2")

        doc.save(str(docx_path))

    def _set_paragraph_border(self, paragraph, border_type="bottom", color="4F81BD", size="6", space="1"):
        """Helper to set a border on a paragraph using low-level XML."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        # Create border element
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)     # Size in 1/8 pt (6 = 0.75pt)
        border.set(qn('w:space'), space) # Spacing in pt
        border.set(qn('w:color'), color)
        
        pBdr.append(border)
        pPr.append(pBdr)

    def _has_visible_border(self, paragraph):
        """Check if paragraph has a border set (used to detect HR lines)."""
        pPr = paragraph._p.find(qn('w:pPr'))
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            return pBdr is not None
        return False        

    def _markdown_to_docx(self, md_path: Path, docx_path: Path, title: str, source_dir: Path | None = None) -> None:
        """Convert Markdown to DOCX via system pandoc."""
        self._require_pandoc()
        
        # 1. Resource path starts with the intermediate MD location
        resource_paths = [str(md_path.parent.resolve())]
        
        # 2. Add the original source directory (where images are)
        if source_dir:
            resource_paths.append(str(source_dir.resolve()))
            
        resource_path_arg = ":".join(resource_paths)

        cmd = [
            "pandoc",
            str(md_path),
            "-o",
            str(docx_path),
            "--standalone",
            f"--resource-path={resource_path_arg}",
            "-f", "markdown+autolink_bare_uris",
        ]
        
        if self.include_toc:
            cmd.append("--toc")
            cmd.append("--toc-depth=3")

        if self.template:
            cmd.append(f"--reference-doc={self.template}")
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise ConversionError(
                f"pandoc failed (exit {result.returncode}): {result.stderr.strip()}"
            )

        # 1. Apply Layout FIRST (Sets Margins, Headers, Footers)
        try:
            self._add_header_footer(docx_path, title)
        except Exception as e:
            print(f"Warning: Failed to add layout/header: {e}")

        # 2. Apply Styling SECOND (Resizes Images/Tables to match new margins)
        try:
            self._apply_styling(docx_path)
        except Exception as e:
            print(f"Warning: Failed to apply styling: {e}")

    def _process_notebook(self, notebook_path: Path, base_name: str, output_dir: Path) -> ConversionResult:
        self._require_nb()
        
        title = self.doc_title if self.doc_title else base_name

        out_html = output_dir / f"{base_name}_report.html"
        out_md = output_dir / f"{base_name}_report.md"
        assets_dir = f"{base_name}_report_files"
        out_docx = output_dir / f"{base_name}_report.docx"

        nb = nbf.read(str(notebook_path), as_version=4)
        nb_clean = self._clean_outputs(nb, keep_text=self.keep_text)

        # Extract DataFrame HTML tables from code cell outputs, convert to
        # markdown pipe tables, and inject as markdown cells so that Pandoc
        # produces real Word tables instead of plain text blocks.
        new_cells = []
        for cell in nb_clean.cells:
            table_htmls = []
            if cell.get("cell_type") == "code":
                remaining = []
                for out in cell.get("outputs", []):
                    data = out.get("data", {})
                    if "text/html" in data and "<table" in data["text/html"]:
                        table_htmls.append(data["text/html"])
                    else:
                        remaining.append(out)
                cell["outputs"] = remaining
            new_cells.append(cell)
            for html in table_htmls:
                md_table = self._html_table_to_markdown(html)
                if md_table:
                    new_cells.append(nbf.v4.new_markdown_cell(source=md_table))
        nb_clean.cells = new_cells

        result = ConversionResult()

        if self.export_html:
            self._export_html(nb_clean, out_html)
            result.html = out_html

        # Export to Markdown (Standard logic, no manual TOC injection)
        self._require_nb()
        c = Config()
        c.MarkdownExporter.exclude_input_prompt = True
        c.MarkdownExporter.exclude_output_prompt = True
        if self.hide_code: c.MarkdownExporter.exclude_input = True
        
        resources = {"output_files_dir": assets_dir}
        md_exporter = MarkdownExporter(config=c)
        md_body, resources = md_exporter.from_notebook_node(nb_clean, resources=resources)

        # Apply Strip Captions
        md_body = self._strip_captions(md_body)
        # Apply HR Removal
        md_body = self._strip_hrs_after_headers(md_body)
        # Apply Figure Styling
        md_body = self._style_captions_in_markdown(md_body)

        out_md.write_text(md_body, encoding="utf-8")
        for fname, data in resources.get("outputs", {}).items():
            img_path = out_md.parent / fname
            img_path.parent.mkdir(parents=True, exist_ok=True)
            img_path.write_bytes(data)

        if self.export_markdown:
            result.markdown = out_md

        self._markdown_to_docx(out_md, out_docx, title=title, source_dir=notebook_path.parent)
        result.docx = out_docx

        if not self.export_markdown:
            out_md.unlink(missing_ok=True)
            assets_path = output_dir / assets_dir
            if assets_path.exists(): shutil.rmtree(assets_path)

        return result

    def _process_markdown(self, markdown_path: Path, base_name: str, output_dir: Path) -> ConversionResult:
        title = self.doc_title if self.doc_title else base_name
        out_docx = output_dir / f"{base_name}.docx"
        
        # Pass the markdown's parent directory as source_dir
        self._markdown_to_docx(markdown_path, out_docx, title=title, source_dir=markdown_path.parent)

        # Read content
        md_content = markdown_path.read_text(encoding="utf-8")
        
        # STRIP CAPTIONS
        md_content = self._strip_captions(md_content)
        # Strip HRs
        md_content = self._strip_hrs_after_headers(md_content)
        # Style Figures
        md_content = self._style_captions_in_markdown(md_content)
        
        # Write to temp or overwrite
        temp_md = output_dir / f"{base_name}_temp.md"
        temp_md.write_text(md_content, encoding="utf-8")
        
        self._markdown_to_docx(temp_md, out_docx, title=title, source_dir=markdown_path.parent)
        temp_md.unlink()
        
        return ConversionResult(docx=out_docx)
    
    def _generate_toc_markdown(self, md_content: str) -> str:
        """Generates a Markdown list of links based on headers."""
        import re
        lines = md_content.split('\n')
        toc_lines = ["## Table of Contents\n"]
        
        # Regex to find Markdown headers: # Header, ## Header, etc.
        header_pattern = re.compile(r'^(#{1,3})\s+(.*)')
        
        for line in lines:
            match = header_pattern.match(line)
            if match:
                level = len(match.group(1))
                title = match.group(2).strip()
                
                # Create anchor link (Pandoc auto-generates these IDs)
                # Pandoc removes punctuation, lowercases, and replaces spaces with hyphens
                anchor = re.sub(r'[^\w\s-]', '', title).lower().replace(' ', '-')
                
                indent = "  " * (level - 1)
                toc_lines.append(f"{indent}- [{title}](#{anchor})")
        
        if len(toc_lines) > 1:
            return "\n".join(toc_lines) + "\n\n"
        return ""
